import json
import subprocess
import sys
import random
import requests
import tools
import csv
import random
import string
import os

from config import data
from docx import Document
from PySide6.QtWidgets import QApplication, QMainWindow, QTableView, QLineEdit, QWidget, QPushButton, QVBoxLayout, QFileDialog
from PySide6.QtCore import QRect
from PySide6.QtGui import QStandardItemModel, QStandardItem, QPixmap
from PySide6.QtCore import QCoreApplication
from add_worker_ui import Ui_Add_Worker
from delete_worker_ui import Ui_Delete_Worker
from add_worker_ui import Ui_Add_Worker
from delete_worker_ui import Ui_Delete_Worker
from not_permission_ui import Ui_NotPermisson
from change_worker_ui import Ui_ChangeWorker
from menu_ui import Ui_MainWindow
from change_menu_ui import Ui_ChangeMenu
from add_card_ui import Ui_Add_card
from add_crypto_ui import Ui_AddCrypto
from add_investor_ui import Ui_AddInvestor
from connect_api_ui import Ui_connect_api
from end_session_ui import Ui_End_session_Window
from start_session_ui import Ui_Start_Session_Window
from delete_invest_ui import Ui_DeleteInvest
from delete_card_ui import Ui_DeleteCard
from delete_crypto_ui import Ui_DeleteCrypto
from buy_and_sell_ui import Ui_Buy_and_Sell
from login_ui import Ui_Login_Form 
from not_permission_ui import Ui_NotPermisson
from parser_ui import Ui_Parser
from net_ui import Ui_Net
from db import Database
from parsers.advertises_config import bybit_data
from parsers.p2p import P2PTrade1
from other_orders_ui import Ui_Other_Order
from tables_ui import Ui_TableForm
from change_crypto_ui import Ui_CryptoChange
from change_card_ui import Ui_ChangeCard
from change_investor_ui import Ui_ChangeInvestor

db = Database('DataAnd.db')

class MainMenu(QMainWindow):
    def __init__(self):
        super(MainMenu, self).__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        self.initialization()
        self.connectionButtons()
    
    def connectionButtons(self): 
        self.tableform = None

        self.ui.btn_translate.clicked.connect(lambda:self.money_transfer()) #кнопка для перевода средств между картами
        self.ui.btn_translate.clicked.connect(lambda:self.set_statistics()) #кнопка для перевода одних средств из одной карты на другое устанавливает другое значение
        self.ui.btn_translate_2.clicked.connect(lambda:StartSession().show()) if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' or db.get_role() == 'ROLE_OWNER' else self.ui.btn_translate_2.clicked.connect(lambda:NotPermission().show()) #Открыть сессию
        self.ui.spred_btn.clicked.connect(lambda: BuyAndSell().show()) #открытие окна покупки и продажи
        self.ui.spred_btn_2.clicked.connect(lambda: self.open_form()) #открытие таблиц
        self.ui.spred_btn_3.clicked.connect(lambda: ParserForm().show()) #кнопка для открытия парсера. 
        
        self.ui.choose_card_box.currentIndexChanged.connect(lambda:self.set_statistics()) #при нажати на комбобокс устанавливаются другие данные
        self.ui.comboBox_9.currentIndexChanged.connect(lambda:self.set_statistics()) #при нажати на комбобокс устанавливаются другие данные
        self.ui.comboBox_8.currentIndexChanged.connect(lambda:self.set_statistics()) #при нажати на комбобокс устанавливаются другие данные
        self.ui.balance_crypto_box.currentIndexChanged.connect(lambda:self.set_statistics()) #при нажати на комбобокс устанавливаются другие данные

        self.ui.balance_crypto_box.currentIndexChanged.connect(lambda:self.set_statistics()) #при нажатии Combobox при выбора криптовалюты обновляются данные

        self.ui.action_8.triggered.connect(lambda:StartSession().show()) if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' or db.get_role() == 'ROLE_OWNER' else self.ui.action_8.triggered.connect(lambda:NotPermission().show()) #Открыть сессию #Открыть сессию
        self.ui.action_11.triggered.connect(lambda:EndSession().show ()) if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' or db.get_role() == 'ROLE_OWNER' else self.ui.action_11.triggered.connect(lambda:NotPermission().show()) #Открыть сессию #Закрыть сессию
        self.ui.action_9.triggered.connect(lambda:db.delete_all_data('history')) if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' or db.get_role() == 'ROLE_OWNER' else self.ui.action_9.triggered.connect(lambda:NotPermission().show()) #Открыть сессию #Очистить ордера
        self.ui.action.triggered.connect(lambda:AddCards().show()) if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' or db.get_role() == 'ROLE_OWNER' else self.ui.action.triggered.connect(lambda:NotPermission().show()) #Открыть сессию #Добавить карту
        self.ui.action_2.triggered.connect(lambda:AddCrypto().show()) if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' else self.ui.action_2.triggered.connect(lambda:NotPermission().show()) #Открыть сессию #Добавить крипту
        self.ui.action_3.triggered.connect(lambda:AddInvestor().show()) if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' else self.ui.action_3.triggered.connect(lambda:NotPermission().show()) #Открыть сессию #Добавляет инвестора
        self.ui.action_API.triggered.connect(lambda:ConnectApi().show()) if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' else self.ui.action_API.triggered.connect(lambda:NotPermission().show()) #Открыть сессию #меню для работы с API
        self.ui.change_crypto_btn.clicked.connect(lambda:ChangeMenu().show()) #запускает меню для обмена крипты
        self.ui.action_19.triggered.connect(lambda:O_Nas_Form().show()) #запускает форму О программе.

        self.ui.action_5.triggered.connect(lambda:DeleteCard().show()) if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' else self.ui.action_5.triggered.connect(lambda: NotPermission().show()) #открывает окно удаление карты
        self.ui.action_6.triggered.connect(lambda:DeleteCrypto().show()) if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' else self.ui.action_6.triggered.connect(lambda: NotPermission().show()) #открывает окно удаление криптовалюты
        self.ui.action_7.triggered.connect(lambda:DeleteInvestor().show()) if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' else self.ui.action_7.triggered.connect(lambda: NotPermission().show()) #открывает окно удаление инвестора
        
        self.ui.action_16.triggered.connect(lambda:self.download_csv('history')) if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' else self.ui.action_16.triggered.connect(lambda: NotPermission().show()) #скачивание таблицы истории
        self.ui.action_17.triggered.connect(lambda:self.download_csv('cards')) if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' else self.ui.action_17.triggered.connect(lambda: NotPermission().show()) #скачивание таблицы карт
        self.ui.action_13.triggered.connect(lambda:self.download_csv('investors')) if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' else self.ui.action_13.triggered.connect(lambda: NotPermission().show()) #скачивание таблицы инвесторов
        
    
    def open_form(self):
        if self.tableform is None:
            self.tableform = TableForm()
        
        self.tableform.show()


    def initialization(self): #инициализация данных. 
#---------Читаем БД------------№        
        cards = db.get_all_elements('cards', 'number_card')
        cryptos = db.get_all_elements('crypto', 'name_crypto')

#---------Стираем данные с форм------------№
        self.ui.balance_crypto_box.clear()

#---------Добавляем карты------------№
        self.ui.choose_card_box.addItems(cards)
        self.ui.comboBox_8.addItems(cards)
        self.ui.comboBox_9.addItems(cards)

#---------Добавляем крипту------------№
        self.ui.balance_crypto_box.addItems(cryptos)

        self.set_statistics()

        self.set_investors()

#---------Основной функционал---------#

    def money_transfer(self):
        card_forward = self.ui.comboBox_9.currentText()
        card_get = self.ui.comboBox_8.currentText()

        name_bank_forward = db.find_element('cards', ['number_card', card_forward], 'name_bank')
        name_bank_get = db.find_element('cards', ['number_card', card_get], 'name_bank') #tools.find_name_bank(card_get)[0]
        bank_sbp_forward = db.find_element('cards', ['number_card', card_forward], 'turnover_sbp')

        if card_get != card_forward:
            summ = float(self.ui.lineEdit_3.text())
            if name_bank_forward != name_bank_get:
                summ = tools.return_percent(summ, 0.5) if float(bank_sbp_forward) - summ <= 0 else summ
                db.set_balance(card_forward, float(summ), 'minus')
                db.set_balance(card_get, float(summ), 'plus')

                db.set_balance(card_forward, float(summ), 'sbp')

                db.set_balance(card_forward, float(summ), 'turnover')
                db.set_balance(card_get, float(summ), 'turnover')
                
            else:
                db.set_balance(card_forward, float(summ), 'minus')
                db.set_balance(card_get, float(summ), 'plus')

                db.set_balance(card_forward, float(summ), 'turnover')
                db.set_balance(card_get, float(summ), 'turnover')
    
    def set_statistics(self):
        #обновление данных при нажатии на comboBox. 
        cards = db.get_all_elements('cards', 'number_card')
        cryptos = db.get_all_elements('crypto', 'name_crypto')

        for card in cards:
            if card not in [self.ui.choose_card_box.itemText(i) for i in range(self.ui.choose_card_box.count())]:
                self.ui.choose_card_box.addItem(card)
                self.ui.comboBox_8.addItem(card)
                self.ui.comboBox_9.addItem(card)
        
        for crypto in cryptos:
            if crypto not in [self.ui.balance_crypto_box.itemText(i) for i in range(self.ui.balance_crypto_box.count())]:
                self.ui.balance_crypto_box.addItem(crypto)

        card_name = self.ui.choose_card_box.currentText()
        card_turnover = db.find_element('cards', ['number_card', card_name], 'turnover')
        card_balance = db.find_element('cards', ['number_card', card_name], 'balance')
        card_turnover_SBP = db.find_element('cards', ['number_card', card_name], 'turnover_sbp')

        crypto_data = self.ui.balance_crypto_box.currentText()
        crypto_quantity = db.find_element('crypto', ['name_crypto', crypto_data], 'quantity')

        self.ui.card_turnover_num.setText(f'{round(card_turnover, 2)}р')
        self.ui.balance_card_num.setText(f'{round(card_balance, 2)}р')
        self.ui.turnover_SBP_num.setText(f'{round(card_turnover_SBP, 2)}р') if float(card_turnover_SBP) >= 0 else self.ui.turnover_SBP_num.setText(f'0р')
        if db.get_from_config('isWorkAPI'):
            try: self.ui.balance_crypto_num.setText(f'{tools.get_balance(crypto_data, "SPOT")["result"]["balance"][0]["transferBalance"]} {crypto_data}') 
            except: self.ui.balance_crypto_num.setText(f'Нет данных!')
        else:
            self.ui.balance_crypto_num.setText(f'{round(float(crypto_quantity), 5)} {crypto_data}')

        cards = db.get_all_elements('cards', 'balance')
        balance = 0
        for item in cards:
            balance += float(item)

        set_balance = db.get_from_config('balance')
        self.ui.all_balance_fiat_num.setText(f'{round(balance, 2)}Р')
        self.ui.payments_to_investors_num.setText('0')
        self.ui.current_income_num.setText(f'{round(balance - set_balance, 2)}') if balance - set_balance >= 0 else self.ui.current_income_num.setText(f'0')

    def set_investors(self):
        investors_summ = db.get_all_elements('investors', 'summ')
        investors_income_percent = db.get_all_elements('investors', 'income')

        all_summ = 0

        for i in range(len(investors_summ)):
            all_summ += (investors_income_percent[i] / 100) * investors_summ[i]
        
        self.ui.payments_to_investors_num.setText(str(round(all_summ, 2)))

    def download_csv(self, table):
        #_______история ордеров________#
        options = QFileDialog.Options()
        file_path, _ =  QFileDialog.getSaveFileName(self, "Сохранить файл", "", "CSV Files (*.csv);;All Files (*)", options=options)
        tab = db.get_table(table)

        with open(file_path + '.csv', mode='w') as file:
            writer = csv.writer(file)
            writer.writerow(tab[1])
            for row in tab[0]:
                writer.writerow(row)

        print('Файл был успешно сохранён!')

class BuyAndSell(QMainWindow):
    def __init__(self):
        super(BuyAndSell, self).__init__()
        self.ui = Ui_Buy_and_Sell()
        self.ui.setupUi(self)
        
        self.initialization()
        self.connectionButtons()

    def initialization(self):
        cards = db.get_all_elements('cards', 'number_card')
        cryptos = db.get_all_elements('crypto', 'name_crypto')
        #---------Стираем данные с форм------------№
        self.ui.card_for_buy_box.clear()
        self.ui.card_for_sell_box.clear()

#---------Добавляем карты------------№
        self.ui.card_for_buy_box.addItems(cards)
        self.ui.card_for_sell_box.addItems(cards)

#---------Добавляем карты------------№
        self.ui.crypto_buy_box.addItems(cryptos)
        self.ui.crypto_sell_box.addItems(cryptos)

    def connectionButtons(self):
        self.ui.btn_buy.clicked.connect(lambda:self.write_data('buy'))
        self.ui.btn_buy.clicked.connect(lambda:self.set_statistics())
        self.ui.btn_buy.clicked.connect(lambda:self.addHistory('buy')) #Добавляет всё в историю если купили
        self.ui.btn_buy.clicked.connect(lambda:self.write_statistic_bank())

        self.ui.btn_sell.clicked.connect(lambda:self.write_data('sell')) #кнопка для продажи криптовалюты
        self.ui.btn_sell.clicked.connect(lambda:self.set_statistics())
        self.ui.btn_sell.clicked.connect(lambda:self.addHistory('sell')) #Добавляет всё в историю если продали
        self.ui.btn_sell.clicked.connect(lambda:self.write_statistic_bank())

        self.ui.fresh_btn.clicked.connect(lambda:self.update_spred(False)) #кнопка для обновления спреда между разными криптопарами
        
        self.ui.enter_price_on_buy.textChanged.connect(lambda:self.update_spred(True)) #строка, которая обновляет данные о спреде, если криптопара одна и та же

        self.ui.enter_price_on_sell.textChanged.connect(lambda:self.update_spred(True)) #строка, которая обновляет данные о спреде, если криптопара одна и та же

    def initialization(self):
        cards = db.get_all_elements('cards', 'number_card')
        cryptos = db.get_all_elements('crypto', 'name_crypto')
        #---------Стираем данные с форм------------№
        self.ui.card_for_buy_box.clear()
        self.ui.card_for_sell_box.clear()
        self.ui.crypto_buy_box.clear()
        self.ui.crypto_sell_box.clear()

        #---------Добавляем карты------------№
        self.ui.card_for_buy_box.addItems(cards)
        self.ui.card_for_sell_box.addItems(cards)

        self.ui.crypto_buy_box.addItems(cryptos)
        self.ui.crypto_sell_box.addItems(cryptos)

    def write_data(self, meaning):
        if tools.contains_only_digits(self.ui.enter_price_on_buy.text()) or tools.contains_only_digits(self.enter_price_on_sell.text()):
            if meaning == 'buy':
                crypto = self.ui.crypto_buy_box.currentText()
                use_SBP_buy_check = self.ui.use_SBP_buy_check.isChecked()

                card = self.ui.card_for_buy_box.currentText()
                
                price = float(self.ui.enter_price_on_buy.text())
                summ = float(self.ui.summ_on_buy_enter.text())
                quantity = summ / price

                db.set_balance(card, float(summ), 'minus')
                db.set_crypto(crypto, float(quantity), 'plus')
                db.set_balance(card, float(summ), 'turnover')

                if use_SBP_buy_check:
                    db.set_balance(card, float(summ), 'sbp')

            elif meaning == 'sell':
                crypto = self.ui.crypto_sell_box.currentText()
                card = self.ui.card_for_sell_box.currentText()
                price = float(self.ui.enter_price_on_sell.text())
                quantity = float(self.ui.summ_on_sell_enter.text())
                
                summ = price * quantity

                db.set_balance(card, float(summ), 'plus')
                db.set_crypto(crypto, float(quantity), 'minus')
                db.set_balance(card, float(summ), 'turnover')

    def update_spred(self, meaning):
        try:
            crypto_buy = self.ui.crypto_buy_box.currentText()
            crypto_sell = self.ui.crypto_sell_box.currentText()
            price_buy = float(self.ui.enter_price_on_buy.text())
            price_sell = float(self.ui.enter_price_on_sell.text())
            if meaning:
                if crypto_buy == crypto_sell:
                    spred = ((price_sell * 100) / price_buy) - 100
                    self.ui.label_3.setText(f'{round(spred, 2)}%')
            else:
                if crypto_buy != crypto_sell:
                    curse, sym = tools.combo_pair(crypto_buy, crypto_sell)
                    if sym:
                        quantity = 1 * curse
                    else:
                        quantity = 1 / curse

                    res = quantity * price_sell
                    spred = ((res * 100) / price_buy) - 100
                    self.ui.label_3.setText(f'{round(spred, 1)}%')
        except Exception as ex: 
            print(ex)

    def change_crypto(self):
        crypto_name_buy = self.ui.crypto_buy_box.currentText()
        crypto_name_sell = self.ui.crypto_sell_box.currentText()

        crypto_quantity_forward = float(self.ui.enter_change_line.text())

        comission = db.get_from_config('commission')
        crypto_quantity = tools.return_percent(crypto_quantity_forward, comission)

        if crypto_name_buy != crypto_name_sell:
            curse, sym = tools.combo_pair(crypto_name_buy, crypto_name_sell)
            if sym:
                quantity = crypto_quantity * curse
            else:
                quantity = crypto_quantity / curse

            db.set_crypto(crypto_name_buy, quantity, 'minus')
            db.set_crypto(crypto_name_sell, quantity, 'plus')

    def set_statistics(self): #для статистики
        #обновление данных при нажатии на comboBox. 
        cards = db.get_all_elements('cards', 'number_card')
        cryptos = db.get_all_elements('crypto', 'name_crypto')

        for card in cards:
            if card not in [self.ui.card_for_buy_box.itemText(i) for i in range(self.ui.card_for_buy_box.count())]:
                self.ui.card_for_buy_box.addItem(card)
                self.ui.card_for_sell_box.addItem(card)
        
        for crypto in cryptos:
            if crypto not in [self.ui.crypto_buy_box.itemText(i) for i in range(self.ui.crypto_buy_box.count())]:
                self.ui.crypto_buy_box.addItem(crypto)
                self.ui.crypto_sell_box.addItem(crypto)

    def addHistory(self, meaning):
        try:
            type = 'BUY' if meaning == 'buy' else 'SELL'
            crypto_name = self.ui.crypto_buy_box.currentText() if meaning == 'buy' else self.ui.crypto_sell_box.currentText()
            summ_fiat = self.ui.summ_on_buy_enter.text() if meaning == 'buy' else f'{float(self.ui.summ_on_sell_enter.text()) * float(self.ui.enter_price_on_sell.text())}'
            price = self.ui.enter_price_on_buy.text() if meaning == 'buy' else self.ui.enter_price_on_sell.text()
            quantity_crypto = f'{float(self.ui.summ_on_buy_enter.text()) / float(self.ui.enter_price_on_buy.text())}' if meaning == 'buy' else self.ui.enter_price_on_buy.text()
            bank = db.find_element('cards', ['number_card', self.ui.card_for_buy_box.currentText()], 'name_bank') if meaning == 'buy' else db.find_element('cards', ['number_card', self.ui.card_for_sell_box.currentText()], 'name_bank')
            sbp = True if self.ui.use_SBP_buy_check.isChecked() and meaning == 'buy' else False
            third_face = True if self.ui.third_party_check.isChecked() and meaning == 'sell' else False
            number_card = self.ui.card_for_buy_box.currentText() if meaning == 'buy' else self.ui.card_for_sell_box.currentText()

            db.add_history(type, crypto_name, summ_fiat, price, quantity_crypto, bank, number_card, sbp, third_face)
        except Exception as ex:
            print('Ошибка - ', ex)

    def write_statistic_bank(self):
        list_status = db.get_all_elements('history', 'status')
        db_list_sbp = db.get_all_elements('history', 'is_sbp')
        db_list_third_face = db.get_all_elements('history', 'is_third_face')
        list_summ_fiat = db.get_all_elements('history', 'summ_fiat')

        quantity_list_buy = []
        quantity_list_sell = []
        medium_price_buy_list = []
        medium_price_sell_list = []
        list_sbp = []
        list_third_face = []

        [quantity_list_buy.append(item) if item == 'BUY' else quantity_list_sell.append(item) for item in list_status]
        [medium_price_buy_list.append(item) if item == 'BUY' else medium_price_sell_list.append(item) for item in list_summ_fiat]
        [list_sbp.append(item) if item == True else None for item in db_list_sbp]
        [list_third_face.append(item) if item == True else None for item in db_list_third_face]

class ChangeMenu(QMainWindow):
    def __init__(self):
        super(ChangeMenu, self).__init__()
        self.ui = Ui_ChangeMenu()
        self.ui.setupUi(self)

        self.initialization()
        self.connectionButtons()
    
    def connectionButtons(self):
        self.ui.pushButton.clicked.connect(lambda: self.change_crypto())
        #self.ui.lineEdit.textChanged.connect(lambda:self.set_price()) #строка, которая обновляет данные о спреде, если криптопара одна и та же

    def initialization(self):
        cryptos = db.get_all_elements('crypto', 'name_crypto')

        self.ui.from_crypto.addItems(cryptos)
        self.ui.to_crypto.addItems(cryptos)

    def change_crypto(self):
        from_crypto = self.ui.from_crypto.currentText()
        to_crypto = self.ui.to_crypto.currentText()

        if tools.contains_only_digits(self.ui.lineEdit.text()) and from_crypto != to_crypto:
            res = tools.change_money(from_crypto, to_crypto, self.ui.lineEdit.text())
            print(res['retMsg'])
            try:
                if res['retMsg'] == "Insufficient balance.":
                    self.ui.label.setText("Недостаточно средств.")
                elif res['retMsg'] == "Order value exceeded lower limit.":
                    self.ui.label.setText("Маленькая заявка.")
                elif res['retMsg'] == "OK":
                    self.ui.label.setText("Заявка прошла успешно!")
                else:
                    self.ui.label.setText("Ошибка!")
            except:
                pass

class AddCards(QMainWindow):
    def __init__(self):
        super(AddCards, self).__init__()
        self.ui = Ui_Add_card()
        self.ui.setupUi(self)
        
        self.connectionButtons()

    def connectionButtons(self):
        self.ui.buttonBox.accepted.connect(lambda:self.accept())
        self.ui.buttonBox.rejected.connect(lambda:self.close())

    def accept(self):
        if self.ui.name_bank_edit.text() != "" and self.ui.number_card_edit.text() != "":
            db.add_card(self.ui.name_bank_edit.text(), self.ui.number_card_edit.text())
            print(self.ui.name_bank_edit.text(), self.ui.number_card_edit.text())
            self.ui.name_bank_edit.setText('')
            self.ui.number_card_edit.setText('')

class AddCrypto(QMainWindow):
    def __init__(self):
        super(AddCrypto, self).__init__()
        self.ui = Ui_AddCrypto()
        self.ui.setupUi(self)

        self.connectionButtons() if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' else NotPermission().show()
    
    def connectionButtons(self):
        self.ui.buttonBox.accepted.connect(lambda:self.addCrypto())
        self.ui.buttonBox.rejected.connect(lambda:self.close())

    def addCrypto(self):
        db.set_crypto(self.ui.name_bank_edit.text(), 0, 'add') if self.ui.name_bank_edit.text() != '' or self.ui.name_bank_edit.text() != None else None
        self.ui.name_bank_edit.setText('')

class AddInvestor(QMainWindow): #тут добавь функционал
    def __init__(self):
        super(AddInvestor, self).__init__()
        self.ui = Ui_AddInvestor()
        self.ui.setupUi(self)
        
        self.connectionButtons() if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' else NotPermission().show()

    def connectionButtons(self):
        self.ui.buttonBox.accepted.connect(lambda:self.addInvestor())
        self.ui.buttonBox.rejected.connect(lambda:self.close())
        
    def addInvestor(self):
        fio = self.ui.lineEdit_2.text()
        summ = self.ui.lineEdit.text()
        income = self.ui.lineEdit_3.text()

        if fio != '' or fio != None and summ != '' or summ != None and income != '' or income != None: 
            db.add_investor(fio, summ, income) 
            
            self.ui.lineEdit_2.clear()
            self.ui.lineEdit_3.clear()
            self.ui.lineEdit.clear()
        else:
            pass

class ConnectApi(QMainWindow):
    def __init__(self):
        super(ConnectApi, self).__init__()
        self.ui = Ui_connect_api()
        self.ui.setupUi(self)

        self.initialization()
        self.connectionButtons() if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' else NotPermission().show()

    def initialization(self): 
        line_api_secret = len(db.get_from_config('api_secret'))
        self.ui.lineEdit.setText(f"{db.get_from_config('api_key')}") if db.get_from_config('api_key') != None else None
        self.ui.lineEdit_2.setText("*" * line_api_secret) if line_api_secret > 0 else None
        self.ui.checkBox.setChecked(db.get_from_config('isWorkAPI'))

    def connectionButtons(self):
        self.ui.buttonBox.accepted.connect(lambda:self.accept())
        self.ui.buttonBox.rejected.connect(lambda:self.close())
        self.ui.checkBox.stateChanged.connect(lambda:db.write_in_config(['isWorkAPI', True]) if self.ui.checkBox.isChecked() else db.write_in_config(['isWorkAPI', False]))
    
    def accept(self):
        db.write_in_config(['api_key', self.ui.lineEdit.text()])
        db.write_in_config(['api_secret', self.ui.lineEdit_2.text()])
        self.close()

class EndSession(QMainWindow):
    def __init__(self):
        super(EndSession, self).__init__()
        self.ui = Ui_End_session_Window()
        self.ui.setupUi(self)

        self.initialization()
        self.connectionButtons() if db.get_role() == 'ROLE_ADMIN' or db.get_role() == 'ROLE_OWNER' else NotPermission().show()

    def initialization(self):
        self.start_summ = sum(db.get_all_elements('config', 'balance'))
        self.summ = sum(db.get_all_elements('cards', 'balance')) - self.start_summ
        self.invest_sum, self.invest_percent = db.get_all_elements('investors', 'summ'), db.get_all_elements('investors', 'income')

        self.invest_income = 0
        for i in range(0, len(self.invest_percent) - 1):
            self.invest_income += self.invest_sum[i] * (self.invest_percent[i] / 100)

        self.result = self.summ - self.invest_income

        self.ui.label_3.setText(QCoreApplication.translate("End_session_Window", f"<html><head/><body><p align=\"right\"><span style=\" font-size:20px;\">{round(self.summ, 2)}</span></p></body></html>", None))
        self.ui.label_5.setText(QCoreApplication.translate("End_session_Window", f"<html><head/><body><p align=\"right\"><span style=\" font-size:20px;\">{round(self.invest_income, 2)}</span></p></body></html>", None))
        self.ui.label_6.setText(QCoreApplication.translate("End_session_Window", f"<html><head/><body><p align=\"right\"><span style=\" font-size:20px;\">{round(self.result, 2)}</span></p></body></html>", None))

    def connectionButtons(self):
        self.ui.pushButton.clicked.connect(lambda:self.save_statistic())

    def save_statistic(self):
        document = Document()

        document.add_heading("Конец торговой сессии!", 0)

        document.add_paragraph(f'Стартовый оборот: {self.start_summ}')
        document.add_paragraph(f'Общий оборот: {self.result}')
        document.add_paragraph(f'Финальня сумма: {self.start_summ + self.summ}')
        document.add_paragraph(f'Комиссия инвесторам: {self.invest_income}')
        document.add_paragraph(f'Выручка: {self.result}')

        options = QFileDialog.Options()
        file_path, _ =  QFileDialog.getSaveFileName(self, "Сохранить файл", "", "docx Files (*.docx);;All Files (*)", options=options)

        file_path += '.docx'
        document.save(file_path)

        print('Файл был успешно сохранён!')

        #тут добавь код на сохранение отчёта.
        #self.close()

class StartSession(QMainWindow):
    def __init__(self):
        super(StartSession, self).__init__()
        self.ui = Ui_Start_Session_Window()
        self.ui.setupUi(self)

        self.initialization()
        self.connectionButtons()

    def initialization(self):
        self.ui.choose_card_box.clear()

        for item in db.get_all_elements('cards', 'number_card'):
            self.ui.choose_card_box.addItem(item)
        self.ui.name_bank.setText(QCoreApplication.translate("Start_Session_Windows", f"<html><head/><body><p align=\"center\"><span style=\" font-size:20pt;\">{db.find_element('cards', ['number_card', self.ui.choose_card_box.currentText()], 'name_bank')}</span></p></body></html>"))

    def connectionButtons(self):
        self.ui.buttonBox.accepted.connect(lambda:self.set_config())
        self.ui.buttonBox.rejected.connect(lambda:self.close())
        self.ui.choose_card_box.currentIndexChanged.connect(lambda:self.ui.name_bank.setText(QCoreApplication.translate("Start_Session_Windows", f"<html><head/><body><p align=\"center\"><span style=\" font-size:20pt;\">{db.find_element('cards', ['number_card', self.ui.choose_card_box.currentText()], 'name_bank')}</span></p></body></html>")))
        self.ui.write_balance_card_btn.clicked.connect(lambda:self.write_bank())

    def set_config(self):
        if tools.contains_only_digits(self.ui.lineEdit.text()):
            db.write_in_config(['balance', self.ui.lineEdit.text()]), self.close() if tools.contains_only_digits(self.ui.lineEdit.text()) else print('no')
            self.ui.lineEdit.setText('')

    def write_bank(self):
        balance = self.ui.lineEdit_2.text()
        if balance != "" and tools.contains_only_digits(self.ui.lineEdit_2.text()):
            db.set_balance(self.ui.choose_card_box.currentText(), float(balance), 'set')
            self.ui.lineEdit_2.setText('')

class DeleteInvestor(QMainWindow):
    def __init__(self):
        super(DeleteInvestor, self).__init__()
        self.ui = Ui_DeleteInvest()
        self.ui.setupUi(self)

        self.initialization()
        self.connectionButtons()

    def initialization(self):
        self.ui.comboBox.clear()
        investors = db.get_investors()

        self.ui.comboBox.addItems(investors)

    def connectionButtons(self):
        self.ui.delete_btn.clicked.connect(lambda:self.del_investor())

    def del_investor(self):
        name = self.ui.comboBox.currentText()
        db.delete_investor(name=name)
        self.initialization()

class DeleteCard(QMainWindow):
    def __init__(self):
        super(DeleteCard, self).__init__()
        self.ui = Ui_DeleteCard()
        self.ui.setupUi(self)

        self.initialization()
        self.connectionButtons()

    def initialization(self):
        self.ui.comboBox.clear()
        cards = db.get_all_elements('cards', 'number_card')

        self.ui.comboBox.addItems(cards)

    def connectionButtons(self):
        self.ui.delete_btn.clicked.connect(lambda:self.del_card())

    def del_card(self):
        card = self.ui.comboBox.currentText()
        db.delete_card(card=card)
        self.initialization()

class DeleteCrypto(QMainWindow):
    def __init__(self):
        super(DeleteCrypto, self).__init__()
        self.ui = Ui_DeleteCrypto()
        self.ui.setupUi(self)

        self.initialization()
        self.connectionButtons()

    def initialization(self):
        self.ui.comboBox.clear()
        cryptos = db.get_all_elements('crypto', 'name_crypto')

        self.ui.comboBox.addItems(cryptos)

    def connectionButtons(self):
        self.ui.delete_btn.clicked.connect(lambda:self.del_crypto())

    def del_crypto(self):
        crypto = self.ui.comboBox.currentText()
        db.delete_crypto(crypto=crypto)
        self.initialization()

class TestClass(QMainWindow):
    def __init__(self):
        super(TestClass, self).__init__()
        self.ui = Ui_Delete_Worker()
        self.ui.setupUi(self)

class DeleteWorkerForm(QMainWindow):
    def __init__(self):
        super(DeleteWorkerForm, self).__init__()
        self.ui = Ui_Delete_Worker()
        self.ui.setupUi(self)

        self.connectionButtons()
        self.initialization()

    def connectionButtons(self):
        self.ui.comboBox.currentIndexChanged.connect(lambda: self.show_data())
        self.ui.pushButton.clicked.connect(lambda: self.delete_data())

    def initialization(self):
        id_workers = db.get_all_elements('workers', 'id')
        print(id_workers)

        for item in id_workers:
            self.ui.comboBox.addItem(str(item))

    def show_data(self):
        id_user = self.ui.comboBox.currentText()
 
        self.ui.label.setText(str(db.get_element_worker_id(id_user, 'username')))
        self.ui.label_2.setText(str(db.get_element_worker_id(id_user, 'subname')))
        self.ui.label_3.setText(str(db.get_element_worker_id(id_user, 'job_title')))
        self.ui.label_4.setText(str(db.get_element_worker_id(id_user, 'salary')))
        self.ui.label_5.setText(str(db.get_element_worker_id(id_user, 'pasport')))

    def delete_data(self):
        db.delete_worker_id(self.ui.comboBox.currentText())
        self.ui.comboBox.clear()
        self.initialization()

class AddWorkerForm(QMainWindow):
    def __init__(self):
        super(AddWorkerForm, self).__init__()
        self.ui = Ui_Add_Worker()
        self.ui.setupUi(self)

        self.connectionButtons()
    
    def connectionButtons(self):
        self.ui.add_btn.clicked.connect(lambda: self.write_data())        

    def write_data(self):
        def is_number(num):
            try:
                float(num)
                return True
            except:
                return False
        username = self.ui.name_input.text()
        surname = self.ui.name_input_2.text()
        job_title = self.ui.comboBox.currentText()
        salary = self.ui.name_input_4.text()
        pasport = self.ui.name_input_5.text()
        characters = characters = string.ascii_letters + string.digits + string.punctuation
        secret_key = ''.join(random.choice(characters) for _ in range(24))

        try:
            db.add_worker(username, surname, job_title, float(salary), pasport, secret_key) if username != '' and surname != '' and job_title != '' and salary != '' and is_number(salary) and pasport !='' else None
            
        except:
            pass

        finally:
            self.ui.name_input.clear()
            self.ui.name_input_2.clear()
            self.ui.name_input_4.clear()
            self.ui.name_input_5.clear()

class ChangeWorkerForm(QMainWindow):
    def __init__(self):
        super(ChangeWorkerForm, self).__init__()
        self.ui = Ui_ChangeWorker()
        self.ui.setupUi(self)

        self.initialization()
        self.connectionButtons()

    def initialization(self):
        id_workers = db.get_all_elements('workers', 'id')
        print(id_workers)

        for item in id_workers:
            self.ui.comboBox.addItem(str(item))

        self.show_data()

    def connectionButtons(self):
        self.ui.pushButton.clicked.connect(lambda: self.save_data())
        self.ui.comboBox.currentIndexChanged.connect(lambda: self.show_data())
    
    def show_data(self):
        id_user = self.ui.comboBox.currentText()
 
        self.ui.lineEdit.setText(str(db.get_element_worker_id(id_user, 'username')))
        self.ui.lineEdit_2.setText(str(db.get_element_worker_id(id_user, 'subname')))
        self.ui.lineEdit_3.setText(str(db.get_element_worker_id(id_user, 'job_title')))
        self.ui.lineEdit_4.setText(str(db.get_element_worker_id(id_user, 'salary')))
        self.ui.lineEdit_5.setText(str(db.get_element_worker_id(id_user, 'pasport')))
        self.ui.lineEdit_6.setText(str(db.get_element_worker_id(id_user, 'secret_key')))
        self.ui.lineEdit_7.setText(str(db.get_element_worker_id(id_user, 'user_id')))

    def save_data(self):
        name = self.ui.lineEdit.text()
        surname = self.ui.lineEdit_2.text()
        job_title = self.ui.lineEdit_3.text()
        salary = self.ui.lineEdit_4.text()
        pasport = self.ui.lineEdit_5.text()
        secret_key = self.ui.lineEdit_6.text()
        user_id = self.ui.lineEdit_7.text()

        id = self.ui.comboBox.currentText()

        db.update_worker(id, name, surname, job_title, salary, pasport, secret_key, user_id)

class ChangeCryptoForm(QMainWindow): #форма по смене криптовалюты
    def __init__(self):
        super(ChangeCryptoForm, self).__init__()
        self.ui = Ui_CryptoChange()
        self.ui.setupUi(self)

        self.initialization()
        self.connectionButtons()

    def initialization(self):
        id_workers = db.get_all_elements('crypto', 'id')
        print(id_workers)

        for item in id_workers:
            self.ui.comboBox.addItem(str(item))

        self.show_data()

    def connectionButtons(self):
        self.ui.pushButton.clicked.connect(lambda: self.save_data())
        self.ui.pushButton_2.clicked.connect(lambda: self.update_actual_balance())
        self.ui.comboBox.currentIndexChanged.connect(lambda: self.show_data())
    
    def show_data(self):
        id_user = self.ui.comboBox.currentText()
 
        self.ui.lineEdit.setText(str(db.get_element_id('crypto', id_user, 'name_crypto')))
        self.ui.lineEdit_2.setText(str(db.get_element_id('crypto', id_user, 'quantity')))

    def save_data(self):
        crypto = self.ui.lineEdit.text()
        quantity = self.ui.lineEdit_2.text()

        id = self.ui.comboBox.currentText()

        db.update_crypto(id, crypto, quantity)
    
    def update_actual_balance(self):
        try:
            balance = round(float(tools.get_balance(self.ui.lineEdit.text(), 'SPOT')["result"]["balance"][0]["transferBalance"]), 5)
            self.ui.lineEdit_2.setText(str(balance))
        except Exception as ex:
            print(ex)

class ChangeCardForm(QMainWindow):
    def __init__(self):
        super(ChangeCardForm, self).__init__()
        self.ui = Ui_ChangeCard()
        self.ui.setupUi(self)

        self.initialization()
        self.connectionButtons()

    def initialization(self):
        id_workers = db.get_all_elements('cards', 'id')
        print(id_workers)

        for item in id_workers:
            self.ui.comboBox.addItem(str(item))

        self.show_data()

    def connectionButtons(self):
        self.ui.pushButton.clicked.connect(lambda: self.save_data())
        self.ui.comboBox.currentIndexChanged.connect(lambda: self.show_data())
    
    def show_data(self):
        id_user = self.ui.comboBox.currentText()
 
        self.ui.lineEdit.setText(str(db.get_element_id('cards', id_user, 'name_bank')))
        self.ui.lineEdit_2.setText(str(db.get_element_id('cards', id_user, 'number_card')))
        self.ui.lineEdit_3.setText(str(db.get_element_id('cards', id_user, 'balance')))
        self.ui.lineEdit_4.setText(str(db.get_element_id('cards', id_user, 'turnover')))
        self.ui.lineEdit_5.setText(str(db.get_element_id('cards', id_user, 'turnover_sbp')))

    def save_data(self):
        name_bank = self.ui.lineEdit.text()
        number_card = self.ui.lineEdit_2.text()
        balance = self.ui.lineEdit_3.text()
        turnover = self.ui.lineEdit_4.text()
        turnover_sbp = self.ui.lineEdit_5.text()

        id = self.ui.comboBox.currentText()

        db.update_card(id, name_bank, number_card, balance, turnover, turnover_sbp)

class ChangeInvestForm(QMainWindow): #форма по смене криптовалюты
    def __init__(self):
        super(ChangeInvestForm, self).__init__()
        self.ui = Ui_ChangeInvestor()
        self.ui.setupUi(self)

        self.initialization()
        self.connectionButtons()

    def initialization(self):
        id_workers = db.get_all_elements('investors', 'id')
        print(id_workers)

        for item in id_workers:
            self.ui.comboBox.addItem(str(item))

        self.show_data()

    def connectionButtons(self):
        self.ui.pushButton.clicked.connect(lambda: self.save_data())
        self.ui.comboBox.currentIndexChanged.connect(lambda: self.show_data())
    
    def show_data(self):
        id_user = self.ui.comboBox.currentText()
 
        self.ui.lineEdit.setText(str(db.get_element_id('investors', id_user, 'name')))
        self.ui.lineEdit_2.setText(str(db.get_element_id('investors', id_user, 'summ')))
        self.ui.lineEdit_3.setText(str(db.get_element_id('investors', id_user, 'income')))
        self.ui.lineEdit_4.setText(str(db.get_element_id('investors', id_user, 'user_id')))

    def save_data(self):
        name = self.ui.lineEdit.text()
        summ = self.ui.lineEdit_2.text()
        income = self.ui.lineEdit_3.text()
        user_id = self.ui.lineEdit_4.text()

        id = self.ui.comboBox.currentText()

        db.update_investor(id, name, summ, income, user_id)

class LoginForm(QMainWindow):
    def __init__(self):
        super(LoginForm, self).__init__()
        self.ui = Ui_Login_Form()
        self.ui.setupUi(self)
        self.is_secret_key = False
        
        self.connectionButtons()

    def connectionButtons(self):
        self.ui.pushButton.clicked.connect(lambda: self.is_check())
        self.ui.lineEdit_2.textChanged.connect(lambda: self.clear_text(1))
        self.ui.lineEdit.textChanged.connect(lambda: self.clear_text(2))
        self.ui.lineEdit_3.textChanged.connect(lambda: self.clear_text(3))

        try:
            self.ui.lineEdit_3.textChanged.connect(lambda:self.confirm_code())
        except:
            pass

    def clear_text(self, meaning):
        try: 
            self.ui.label.deleteLater() if meaning == 1 else None
        except:
            pass
        
        try:
            self.ui.label_2.deleteLater() if meaning == 2 else None
        except:
            pass
        
        try: 
            self.ui.label_3.deleteLater() if meaning == 3 else None
        except: 
            pass

    def is_check(self):
        global data
        name = self.ui.lineEdit_2.text()
        pasport = self.ui.lineEdit.text()
        if db.check_login_and_password(name, pasport):
            cod = random.randint(10000, 99999)
            user_id = db.get_user_id(pasport)
            try:
                url = f"https://api.telegram.org/bot{db.get_bot_token()}/sendMessage"
                payload = {
                    'chat_id': user_id,
                    'text': f'Ваш код для подтвержения входа: {cod}'
                }
                response = requests.post(url, data=payload)
                
                data[pasport] = cod  # Сохраняем код для пользователя
                self.ui.pushButton.setText('Повтор')

                self.connectionButtons()

            except Exception as e: 
                self.ui.pushButton.setText('Секретный ключ')
                self.is_secret_key = True
            

    def confirm_code(self):
        enter_code = self.ui.lineEdit_3.text()
        if not self.is_secret_key:
            try:
                if int(enter_code) == data[self.ui.lineEdit.text()]:
                    db.set_work(self.ui.lineEdit.text(), True)
                    self.close()
                    MainMenu().show() 
                    #закрывает логин форму
                else:
                    print("Неверный код")
            except Exception as ex:
                print(ex)
        else:
            name = self.ui.lineEdit_2.text()
            pasport = self.ui.lineEdit.text()
            true_key = db.get_secret_key(name, pasport)
            if str(enter_code) == true_key:
                db.set_work(self.ui.lineEdit.text(), True)
                self.close()
                MainMenu().show() 
                #закрывает логин форму
            else:
                print("Неверный код")

class ParserForm(QMainWindow): 
    def __init__(self): 
        super(ParserForm, self).__init__()
        self.ui = Ui_Parser()
        self.ui.setupUi(self)

        self.initialization()
        self.connectionButtons()

        self.pos = 'taker_maker'
        self.set_best_orders()

    def initialization(self):
        self.ui.comboBox_2.clear()
        self.ui.comboBox_3.clear()

        self.ui.label_6.setText('') 
        self.ui.label_9.setText('') 
        self.ui.label_13.setText('')
        self.ui.label_16.setText('')
        self.ui.label_19.setText('')
        self.ui.label_7.setText('') 
        self.ui.label_10.setText('')
        self.ui.label_12.setText('')
        self.ui.label_17.setText('')
        self.ui.label_20.setText('')
        self.ui.label_29.setText('')
        self.ui.label_30.setText('')
        self.ui.label_31.setText('')
        self.ui.label_32.setText('')
        self.ui.label_33.setText('') 
        
        if self.ui.comboBox.currentText() == 'BYBIT':
            list_crypto_bybit = bybit_data['bybit']['crypto']
            list_bank_bybit = bybit_data['bybit']['bank']
            list_list_bank_bybit = []
            
            for key, value in list_bank_bybit.items():
                list_list_bank_bybit.append(key) 
            
            self.ui.comboBox_2.addItems(list_crypto_bybit)
            self.ui.comboBox_3.addItems(list_list_bank_bybit)

        elif self.ui.comboBox.currentText() == "HUOBI":
            pass
        
        else: 
            pass

    def connectionButtons(self):
        self.ui.comboBox.currentIndexChanged.connect(lambda: self.initialization())
        self.ui.pushButton_6.clicked.connect(lambda: self.update_order(self.ui.comboBox.currentText()))
        self.ui.pushButton_7.clicked.connect(lambda: self.change_button())
        self.ui.pushButton_2.clicked.connect(lambda: OtherOrder().show())
        #self.ui.pushButton_2.clicked.connect(lambda: OtherOrder().show)
        self.ui.pushButton.clicked.connect(lambda: subprocess.Popen(["python", os.path.join('parsers', 'find_best_price.py')]))
        self.ui.pushButton_3.clicked.connect(lambda: self.set_pos('maker_taker'))
        self.ui.pushButton_4.clicked.connect(lambda: self.set_pos('taker_maker'))
        self.ui.pushButton_5.clicked.connect(lambda: self.set_pos('maker_maker'))

    def change_button(self):
        if self.ui.pushButton_7.text() == 'ПОКУПКА':
            self.ui.pushButton_7.setStyleSheet(u"color: white;\n"
    "background: red;\n"
    "border-radius: 10px;\n"
    "font-size: 28px;")
            self.ui.pushButton_7.setText(QCoreApplication.translate("MainWindow", u"\u041f\u0420\u041e\u0414\u0410\u0416\u0410", None))
        elif self.ui.pushButton_7.text() == 'ПРОДАЖА':
            self.ui.pushButton_7.setStyleSheet(u"color: white;\n"
    "background: lime;\n"
    "border-radius: 10px;\n"
    "font-size: 28px;")
            self.ui.pushButton_7.setText(QCoreApplication.translate("MainWindow", u"\u041f\u041e\u041a\u0423\u041f\u041a\u0410", None))

    def set_pos(self, pos):
        self.pos = pos
        if pos == 'maker_taker':
            self.ui.pushButton_3.setStyleSheet(u"color: white;\n"
    "background: #00F0FF;\n"
    "border-radius: 10px;\n"
    "font-size: 28px;")
            self.ui.pushButton_4.setStyleSheet(u"color: white;\n"
    "background: #000AFF;\n"
    "border-radius: 10px;\n"
    "font-size: 28px;")
            self.ui.pushButton_5.setStyleSheet(u"color: white;\n"
    "background: #000AFF;\n"
    "border-radius: 10px;\n"
    "font-size: 28px;")
        elif pos == 'taker_maker':
            self.ui.pushButton_3.setStyleSheet(u"color: white;\n"
    "background: #000AFF;\n"
    "border-radius: 10px;\n"
    "font-size: 28px;")
            self.ui.pushButton_4.setStyleSheet(u"color: white;\n"
    "background: #00F0FF;\n"
    "border-radius: 10px;\n"
    "font-size: 28px;")
            self.ui.pushButton_5.setStyleSheet(u"color: white;\n"
    "background: #000AFF;\n"
    "border-radius: 10px;\n"
    "font-size: 28px;")
        elif pos == 'maker_maker':
            self.ui.pushButton_3.setStyleSheet(u"color: white;\n"
    "background: #000AFF;\n"
    "border-radius: 10px;\n"
    "font-size: 28px;")
            self.ui.pushButton_4.setStyleSheet(u"color: white;\n"
    "background: #000AFF;\n"
    "border-radius: 10px;\n"
    "font-size: 28px;")
            self.ui.pushButton_5.setStyleSheet(u"color: white;\n"
    "background: #00F0FF;\n"
    "border-radius: 10px;\n"
    "font-size: 28px;")

        self.set_best_orders()

    def set_best_orders(self):
        with open(f'parsers/data/{self.pos}.json', 'r', encoding='utf-8') as file:
            data = json.load(file)

        try:
            buy_crypto = data['buy_crypto']
            buy_price = data['buy_price']
            buy_order = data['buy_order']

            sell_crypto = data['sell_crypto']
            sell_price = data['sell_price']
            sell_order = data['sell_order']

            percent = data['percent']

            self.ui.label_23.setText(str(buy_crypto))
            self.ui.label_25.setText(str(buy_price))
            self.ui.label_26.setText(str(buy_order))

            self.ui.label_24.setText(str(sell_crypto))
            self.ui.label_27.setText(str(sell_price))
            self.ui.label_28.setText(str(sell_order))

            self.ui.label_35.setText(str(percent) + '%')

        except Exception as ex:
           self.ui.label_23.setText(str())
           self.ui.label_25.setText("Обнови курсы!")
           self.ui.label_26.setText(str())

           self.ui.label_24.setText(str())
           self.ui.label_27.setText("Обнови курсы!")
           self.ui.label_28.setText(str())

           self.ui.label_35.setText('0%')
           print(ex)

    def update_order(self, exchange):
        crypto = self.ui.comboBox_2.currentText()
        side = 'buy' if self.ui.pushButton_7.text() == 'ПОКУПКА' else 'sell'
        payments = bybit_data[exchange.lower()]['bank'][self.ui.comboBox_3.currentText()]

        cl = P2PTrade1(fiat='RUB', crypto=crypto, side=side, payments=payments)
        if exchange == 'BYBIT': 
            try:
                data = cl.bybitP2P()
                items = data['result']['items']

                try:
                    self.ui.label_6.setText(str(items[0]['price']))
                except (IndexError, KeyError, TypeError):
                    self.ui.label_6.setText('')

                try:
                    self.ui.label_9.setText(str(items[1]['price']))
                except (IndexError, KeyError, TypeError):
                    self.ui.label_9.setText('')

                try:
                    self.ui.label_13.setText(str(items[2]['price']))
                except (IndexError, KeyError, TypeError):
                    self.ui.label_13.setText('')

                try:
                    self.ui.label_16.setText(str(items[3]['price']))
                except (IndexError, KeyError, TypeError):
                    self.ui.label_16.setText('')

                try:
                    self.ui.label_19.setText(str(items[4]['price']))
                except (IndexError, KeyError, TypeError):
                    self.ui.label_19.setText('')

                try:
                    self.ui.label_7.setText(f'https://www.bybit.com/fiat/trade/otc/profile/{items[0]["userId"]}/{crypto}/RUB/item')
                except Exception as ex:
                    self.ui.label_7.setText('')
                    print('Ошибка 1', ex)

                try:
                    self.ui.label_10.setText(f'https://www.bybit.com/fiat/trade/otc/profile/{items[1]["userId"]}/{crypto}/RUB/item')
                except (IndexError, KeyError, TypeError):
                    self.ui.label_10.setText('')

                try:
                    self.ui.label_12.setText(f'https://www.bybit.com/fiat/trade/otc/profile/{items[2]["userId"]}/{crypto}/RUB/item')
                except (IndexError, KeyError, TypeError):
                    self.ui.label_12.setText('')

                try:
                    self.ui.label_17.setText(f'https://www.bybit.com/fiat/trade/otc/profile/{items[3]["userId"]}/{crypto}/RUB/item')
                except (IndexError, KeyError, TypeError):
                    self.ui.label_17.setText('')

                try:
                    self.ui.label_20.setText(f'https://www.bybit.com/fiat/trade/otc/profile/{items[4]["userId"]}/{crypto}/RUB/item')
                except (IndexError, KeyError, TypeError):
                    self.ui.label_20.setText('')

                try:
                    self.ui.label_29.setText(items[0]['nickName'])
                except Exception as ex:
                    self.ui.label_29.setText('')
                    print('Ошибка 2', ex)

                try:
                    self.ui.label_30.setText(items[1]['nickName'])
                except (IndexError, KeyError, TypeError):
                    self.ui.label_30.setText('')

                try:
                    self.ui.label_31.setText(items[2]['nickName'])
                except (IndexError, KeyError, TypeError):
                    self.ui.label_31.setText('')

                try:
                    self.ui.label_32.setText(items[3]['nickName'])
                except (IndexError, KeyError, TypeError):
                    self.ui.label_32.setText('')

                try:
                    self.ui.label_33.setText(items[4]['nickName'])
                except (IndexError, KeyError, TypeError):
                    self.ui.label_33.setText('')
            except Exception as ex:
                print(ex)

class OtherOrder(QMainWindow):
    def __init__(self):
        super(OtherOrder, self).__init__()
        self.ui = Ui_Other_Order()
        self.ui.setupUi(self)

        self.initialization()
    
    def initialization(self):
        try:
            with open('parsers/data/taker_maker.txt', 'r', encoding='utf-8') as file:
                text = file.read()
                self.ui.textEdit.setText(text)
        except:
            pass
        
        try: 
            with open('parsers/data/maker_maker.txt', 'r', encoding='utf-8') as file:
                text = file.read()
                self.ui.textEdit_2.setText(text)
        except:
            pass
        
        try: 
            with open('parsers/data/maker_taker.txt', 'r', encoding='utf-8') as file:
                text = file.read()
                self.ui.textEdit_3.setText(text)
        except: 
            pass

class NotPermission(QMainWindow):
    def __init__(self):
        super(NotPermission, self).__init__()
        self.ui = Ui_NotPermisson()
        self.ui.setupUi(self)
        
        self.ui.pushButton.clicked.connect(lambda: self.close())
        #self.ui.pushButton.clicked.connect(lambda: O_Nas_Form.show()) тут продолжи

class O_Nas_Form(QMainWindow):
    def __init__(self):
        super(O_Nas_Form, self).__init__()
        self.ui = Ui_Net()
        self.ui.setupUi(self)

        self.ui.pushButton.clicked.connect(lambda: self.close())
        self.ui.pushButton_2.clicked.connect(lambda: self.send_message())
        
    def send_message(self):
        url = f"https://api.telegram.org/bot{db.get_bot_token()}/sendMessage"
        payload = {
            'chat_id': 5500790836,
            'text': f'Привет с НЭТ!', 
        }
        response = requests.post(url, data=payload)

class TableForm(QMainWindow):
    def __init__(self):
        super(TableForm, self).__init__()
        self.ui = Ui_TableForm()
        self.ui.setupUi(self)

        self.initialization()
        self.connectionButtons()

    def initialization(self):
        self.view_data('cards', self.ui.tableView_2)
        self.view_data('crypto', self.ui.tableView_3)
        self.view_data('investors', self.ui.tableView)
        self.view_data('workers', self.ui.tableView_4)   

    def connectionButtons(self):
        #добавляем воркеров
        self.ui.pushButton.clicked.connect(lambda: AddWorkerForm().show()) if db.get_role() == 'ROLE_OWNER' else self.ui.pushButton.clicked.connect(lambda: NotPermission().show())
        self.ui.pushButton_2.clicked.connect(lambda: DeleteWorkerForm().show()) if db.get_role() == 'ROLE_OWNER' else self.ui.pushButton_2.clicked.connect(lambda: NotPermission().show())
        self.ui.pushButton_3.clicked.connect(lambda: ChangeWorkerForm().show()) if db.get_role() == 'ROLE_OWNER' else self.ui.pushButton_3.clicked.connect(lambda: NotPermission().show())

        #добавляем карты
        self.ui.pushButton_10.clicked.connect(lambda: AddCards().show()) if db.get_role() == 'ROLE_OWNER' else self.ui.pushButton_10.clicked.connect(lambda: NotPermission().show())
        self.ui.pushButton_11.clicked.connect(lambda: ChangeCardForm().show()) if db.get_role() == 'ROLE_OWNER' else self.ui.pushButton_11.clicked.connect(lambda: NotPermission().show())
        self.ui.pushButton_12.clicked.connect(lambda: DeleteCard().show()) if db.get_role() == 'ROLE_OWNER' else self.ui.pushButton_12.clicked.connect(lambda: NotPermission().show())

        #добавляем крипту
        self.ui.pushButton_8.clicked.connect(lambda: AddCrypto().show()) if db.get_role() == 'ROLE_OWNER' else self.ui.pushButton_8.clicked.connect(lambda: NotPermission().show())
        self.ui.pushButton_9.clicked.connect(lambda: ChangeCryptoForm().show()) if db.get_role() == 'ROLE_OWNER' else self.ui.pushButton_9.clicked.connect(lambda: NotPermission().show())
        self.ui.pushButton_7.clicked.connect(lambda: DeleteCrypto().show()) if db.get_role() == 'ROLE_OWNER' else self.ui.pushButton_7.clicked.connect(lambda: NotPermission().show())

        #добавляем инвесторов
        self.ui.pushButton_5.clicked.connect(lambda: AddInvestor().show()) if db.get_role() == 'ROLE_OWNER' else self.ui.pushButton_5.clicked.connect(lambda: NotPermission().show())
        self.ui.pushButton_6.clicked.connect(lambda: ChangeInvestForm().show()) if db.get_role() == 'ROLE_OWNER' else self.ui.pushButton_6.clicked.connect(lambda: NotPermission().show())
        self.ui.pushButton_4.clicked.connect(lambda: DeleteInvestor().show()) if db.get_role() == 'ROLE_OWNER' else self.ui.pushButton_4.clicked.connect(lambda: NotPermission().show())
        
    def view_data(self, table_name, element):
        data, headers = db.get_table(table_name)
        model = QStandardItemModel(0, len(headers), self)
        model.setHorizontalHeaderLabels(headers)

        for row in data:
            items = [QStandardItem(str(field)) for field in row]
            model.appendRow(items)

        element.setModel(model)
        
if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = TableForm()
    
    window.show()
    try:
        sys.exit(app.exec())
    except SystemExit:
        #db.set_work(None, False)
        pass